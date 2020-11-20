from docx import Document

document = Document()

# section: user basic info.
name = input('Enter your full name: ')
document.add_heading(name)

user_info = document.add_paragraph()

address = input('Enter your address: ')
user_info.add_run(address + '\n').italic = True

email = input('Enter your email: ')
user_info.add_run(email + '\n')

phone_number = input('Enter your phone number: ')
user_info.add_run(phone_number)

# section: user bio.
document.add_heading('Career Profile')
user_bio = input('Describe yourself: ')
document.add_paragraph(user_bio)

# section: skills
document.add_heading('Skills')

while True:
    has_skill = input('Do you have skill: (y/n) ')

    if has_skill.lower() == 'y':
        skill_name = input('Enter your skill name: ')
        document.add_paragraph(skill_name, 'List Bullet')
    else:
        break

# section: experience
has_experience = input('Do you have experience: (y/n) ')

# if experience then add heading.
if has_experience == 'y':
    document.add_heading('Experiences')
    while True:
        has_more_experience = input('Do you have more experience: (y/n) ')

        experience_paragraph = document.add_paragraph()

        if has_more_experience == 'y':
            company_name = input('Enter the company name: ')
            from_date = input('From Date (year): ')
            to_date = input('To Date (year): ')
            experience = input(f'Describe your experience at {company_name}: ')

            experience_paragraph.add_run(company_name + ', ').bold = True
            experience_paragraph.add_run(from_date + '-' + to_date + '\n')
            experience_paragraph.add_run(experience)
        else:
            break

# section: education
document.add_heading('Education')

while True:
    has_education = input('Do you have education: (y/n) ')

    education_paragraph = document.add_paragraph()

    if has_education == 'y':
        education_name = input('Enter education institute name: ')
        from_date = input('From Date (year): ')
        to_date = input('To Date (year): ')

        education_paragraph.add_run(education_name).bold = True
        education_paragraph.add_run('\n' + from_date + '-' + to_date)
    else:
        break

document.save('cv.docx')
